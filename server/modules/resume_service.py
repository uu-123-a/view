"""简历文本提取与讯飞星火结构化分析。"""

from __future__ import annotations

import io
import json
import re
from pathlib import Path

from docx import Document
from pypdf import PdfReader
from werkzeug.datastructures import FileStorage

from .spark_service import SparkService

ALLOWED_SUFFIXES = {".pdf", ".docx", ".txt"}
KNOWN_SKILLS = [
    "Python", "Java", "C++", "PyTorch", "TensorFlow", "Transformer", "RAG",
    "LangChain", "Flask", "Django", "FastAPI", "Vue", "React", "MySQL",
    "Redis", "Docker", "Kubernetes", "Linux", "Git", "NLP", "计算机视觉",
    "多模态", "大模型", "机器学习", "深度学习",
]


class ResumeService:
    def __init__(self) -> None:
        self.spark = SparkService()

    @staticmethod
    def extract(upload: FileStorage) -> tuple[str, str]:
        filename = upload.filename or "resume.txt"
        suffix = Path(filename).suffix.lower()
        if suffix not in ALLOWED_SUFFIXES:
            raise ValueError("仅支持 PDF、DOCX 和 TXT 简历。")
        data = upload.read(5 * 1024 * 1024 + 1)
        if len(data) > 5 * 1024 * 1024:
            raise ValueError("简历文件不能超过 5MB。")
        if not data:
            raise ValueError("上传的简历文件为空。")

        try:
            if suffix == ".pdf":
                reader = PdfReader(io.BytesIO(data))
                text = "\n".join(page.extract_text() or "" for page in reader.pages)
            elif suffix == ".docx":
                document = Document(io.BytesIO(data))
                parts = [paragraph.text for paragraph in document.paragraphs]
                for table in document.tables:
                    parts.extend(" | ".join(cell.text for cell in row.cells) for row in table.rows)
                text = "\n".join(parts)
            else:
                try:
                    text = data.decode("utf-8-sig")
                except UnicodeDecodeError:
                    text = data.decode("gb18030")
        except Exception as exc:
            raise ValueError("无法解析该简历，请确认文件没有损坏或加密。") from exc

        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text).strip()
        if len(text) < 30:
            raise ValueError("简历中可提取的文字太少；扫描版 PDF 请先进行 OCR。")
        return text[:30000], suffix[1:]

    @staticmethod
    def _fallback(content: str) -> dict:
        lowered = content.lower()
        skills = [skill for skill in KNOWN_SKILLS if skill.lower() in lowered][:12]
        return {
            "summary": "已提取简历文本，可直接用于生成岗位相关面试问题。",
            "skills": skills or ["待在面试中进一步确认"],
            "projects": [],
            "risks": ["建议补充项目中的个人贡献、量化结果和技术难点。"],
            "questions": [
                "请介绍简历中最能体现你技术能力的项目。",
                "你在该项目中承担了哪些不可替代的工作？",
                "项目结果可以用哪些数据进行量化？",
            ],
            "source": "fallback",
        }

    @classmethod
    def _normalize(cls, raw: dict, content: str) -> dict:
        fallback = cls._fallback(content)
        def strings(value: object, default: list[str], limit: int) -> list[str]:
            if not isinstance(value, list):
                return default
            result = [str(item).strip()[:160] for item in value if str(item).strip()]
            return result[:limit] or default

        projects: list[dict[str, str]] = []
        if isinstance(raw.get("projects"), list):
            for item in raw["projects"][:5]:
                if isinstance(item, dict):
                    projects.append({
                        "name": str(item.get("name") or "项目经历")[:40],
                        "highlight": str(item.get("highlight") or "待深入追问")[:180],
                    })
                elif str(item).strip():
                    projects.append({"name": "项目经历", "highlight": str(item)[:180]})
        return {
            "summary": str(raw.get("summary") or fallback["summary"])[:300],
            "skills": strings(raw.get("skills"), fallback["skills"], 12),
            "projects": projects,
            "risks": strings(raw.get("risks"), fallback["risks"], 6),
            "questions": strings(raw.get("questions"), fallback["questions"], 8),
            "source": "spark",
        }

    def analyze(self, content: str) -> dict:
        prompt = f"""你是中文技术招聘专家。分析下面的候选人简历，只输出合法 JSON，不要 Markdown。
JSON 格式：
{{"summary":"100字内职业概述","skills":["技能"],"projects":[{{"name":"项目名","highlight":"技术和成果"}}],"risks":["需要追问或补充之处"],"questions":["针对简历的面试问题"]}}
要求：skills 最多12项，projects最多5项，risks最多6项，questions最多8项；不得编造简历中不存在的经历。

简历：
{content[:12000]}"""
        try:
            response = self.spark.chat(
                [{"role": "user", "content": prompt}], temperature=0.2,
                max_tokens=1800, uid="resume-analysis",
            )
            match = re.search(r"\{.*\}", response, re.DOTALL)
            if match is None:
                raise ValueError("星火未返回 JSON")
            raw = json.loads(match.group(0))
            if not isinstance(raw, dict):
                raise ValueError("分析结果不是对象")
            return self._normalize(raw, content)
        except (RuntimeError, ValueError, json.JSONDecodeError, TypeError):
            return self._fallback(content)
